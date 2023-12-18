# 기본 설정
import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import koreanize_matplotlib
import base64
import io
import docx
from docx.shared import Inches
import docx2txt
from PIL import Image

# Streamlit 앱의 제목 설정
st.title("작성된 기사 읽고 비판하기")

# MS Word 파일 업로드하기
uploaded_file = st.file_uploader("Word 파일을 업로드하세요.", type=["docx"])

# 파일이 업로드되었는지 확인
if uploaded_file is not None:
    text = docx2txt.process(uploaded_file)          # word 파일을 텍스트 형태로 읽음
    st.markdown('**<다음은 업로드한 파일의 텍스트 부분과 그림입니다. 기사를 검토하세요.>**')
    st.markdown(f'<div style="padding: 10px; border: 1px solid #e6e6e6">{text}</div>',
                unsafe_allow_html=True)             # 텍스트를 페이지에 표시

    document = docx.Document(uploaded_file)         # word 파일에서 이미지 추출
    images = []

    for rel in document.part.rels.values():
        if "image" in rel.reltype:
            image_data = rel.target_part.blob
            image = Image.open(io.BytesIO(image_data))
            images.append(image)

    # 이미지 출력
    for image in images:
        st.image(image)

st.subheader("기사를 바탕으로 내용을 비판해봅시다.")
article = st.text_area("기사의 내용이 타당한지 적어보세요.", height=200)
student_info = st.text_input("학번과 이름을 입력하세요 (예: 10801 김수훈)")  # 학생 정보 input widget

if st.button("검토 파일 생성하기"):
    if not article or not student_info:
        st.error("내용과 학번/이름을 모두 입력한 후, 파일을 생성하세요.")     # article이 입력되지 않았다면 오류 메시지 출력
    else:
        doc_cr = docx.Document()         # 비어있는 docx 파일 생성
        doc_cr.add_paragraph(article)
        
        # MS Word 파일 생성
        file_data_cr = io.BytesIO()
        doc_cr.save(file_data_cr)
        file_data_cr.seek(0)

        # MS Word 파일 다운로드 링크 생성    
        b64_cr_txt = base64.b64encode(file_data_cr.getvalue()).decode()
        href = f'<a href="data:application/octet-stream;base64,{b64_cr_txt}" download="{student_info}_article_cr.docx">MS Word 파일로 다운로드하기</a>'
        if article and student_info:
            st.markdown(href, unsafe_allow_html=True)
