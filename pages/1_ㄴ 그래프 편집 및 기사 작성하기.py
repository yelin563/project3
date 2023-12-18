# 기본 설정
import streamlit as st
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import koreanize_matplotlib
import base64
import io
import docx
from docx.shared import Inches

#session_state에서 호출
df = st.session_state['df']
col_list = st.session_state['col_list']
row_list = st.session_state['row_list']
graph_option = st.session_state['graph_option']

# 앱에서 사용할 기능 함수 정의 (그래프를 이미지 파일로 저장하는 작업) <-- 메인페이지와 동일
def save_graph_to_image(fig):
    img_data = io.BytesIO()
    fig.savefig(img_data, format='png')
    img_data.seek(0)
    return img_data
def create_download_link(data, filename, text):
    b64_ori = base64.b64encode(data).decode()
    href = f'<a href="data:file/png;base64,{b64_ori}" download="{filename}">{text}</a>'
    return href

# Streamlit 앱의 제목 설정
st.title("기사 작성을 위해 그래프 편집하기")

# 행/열 선택 위젯 생성
st.subheader("전체 통계자료로부터 그래프를 그릴 자료를 추려봅시다.")
option = [#"행",            <-현재 오류로 미구현
       "열"]
selected_opt = st.selectbox("기준: ", option, key="selected_opt")
if selected_opt == "열":
    selected_list = st.selectbox("선택된 열:", col_list, key="selected_col_list")
    selected_df = df[selected_list]
else:           # <- option 미구현 상태로, else 문이 작동하지 않게 되어있음
    selected_list = st.selectbox("선택된 행:", row_list, key="selected_row_list")
    selected_df = df[selected_list]
    
# 데이터로 그래프 그리기
fig_sel, ax=plt.subplots()
graph_selected_opt = st.selectbox("어떤 그래프로 나타낼까요?", graph_option, key="graph_selected_opt")
if graph_selected_opt == "히스토그램":
    df_sel = selected_df.astype(float) # 숫자 데이터를 실수형으로 변환(정수->실수 or 실수->실수)
    value_m = int(min(selected_df.values.flatten()))
    value_M = int(max(selected_df.values.flatten())) # 다시 정수형으로 변환
    bin_size = st.slider("계급의 크기를 선택하세요.", min_value=value_m, max_value=value_M, value=round(value_M/20), step=1) #value:초기값
    selected_df.plot(kind="hist", bins=range(value_m, value_M + bin_size, bin_size), ax=ax)
else:
    if graph_selected_opt == "막대그래프":
        selected_df.plot(kind="bar", ax=ax)
    elif graph_selected_opt == "꺾은선그래프":
        selected_df.plot(kind="line", ax=ax, marker="o")
    elif graph_selected_opt == "원그래프":
        selected_df.plot.pie(y=selected_df.columns[0], ax=ax)
        ax.legend(bbox_to_anchor=(1, 1))
ax.set_xlabel("")
ax.set_ylabel("")
if graph_selected_opt != "그래프 선택하기":
    st.pyplot(fig_sel)
    # 그래프를 이미지로 변환하여 다운로드 링크 생성 (session_state에 저장된 함수 사용)
    graph_image_sel = save_graph_to_image(fig_sel)
    download_link_sel = create_download_link(graph_image_sel.getvalue(), f"편집한 그래프.png", "여기를 눌러 그래프를 다운로드하세요.")
    st.markdown(download_link_sel, unsafe_allow_html=True)
# 데이터에서 통계량을 출력하는 코드
    stats_option = ["통계량을 선택하세요.", "자료의 개수", "평균", "중앙값", "최빈값", "분산", "표준편차"] # 다양한 통계량을 출력할 수 있도록 기능 개선하면 좋음
    stats_selected_opt = st.selectbox("어떤 통계량을 계산할까요?", stats_option, key="stats_selected_opt")
    if stats_selected_opt == "자료의 개수":
        count = selected_df.count()
        st.write(f"자료의 개수는 {count}개 입니다.")
    elif stats_selected_opt == "평균":
        mean = selected_df.mean()
        st.write(f"평균은 {mean} 입니다.")
    elif stats_selected_opt == "중앙값":
        median = selected_df.median()
        st.write(f"중앙값은 {median} 입니다.")
    elif stats_selected_opt == "최빈값":
        mode = selected_df.mode()
        st.write(f"최빈값은 {mode} 입니다.")
    elif stats_selected_opt == "분산":
        variance = selected_df.var()
        st.write(f"분산은 {variance} 입니다.")
    elif stats_selected_opt == "표준편차":
        std = selected_df.std()
        st.write(f"표준편차는 {std} 입니다.")
    
st.subheader("자료를 바탕으로 기사를 작성해봅시다.")
article = st.text_area("기사를 작성하세요.", height=200)

if st.button("기사 파일 생성하기"):
    doc = docx.Document()         # 비어있는 docx 파일 생성
    if graph_selected_opt != "그래프 선택하기":        # 그래프가 그려져 있으면 그래프를 저장하여 doc 파일에 추가
        plt.show()
        plt.savefig('graph.png')
        doc.add_picture('graph.png', width=Inches(6))
    doc.add_paragraph(article)
    
    # MS Word 파일 생성
    file_data = io.BytesIO()
    doc.save(file_data)
    file_data.seek(0)

    # MS Word 파일 다운로드 링크 생성
    b64_txt = base64.b64encode(file_data.getvalue()).decode()
    href = f'<a href="data:application/octet-stream;base64,{b64_txt}" download="article.docx">MS Word 파일로 다운로드하기</a>'
    st.markdown(href, unsafe_allow_html=True)



## 도전과제: 막대/꺾은선 그래프의 경우 y축 중간을 생략하는 코드의 구현(feat. df)